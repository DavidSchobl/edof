"""DOCX (Microsoft Word) import/export for EDOF document-mode files.

v4.1.24.0 — FIRST PASS. Scope and limitations are deliberate:

What round-trips (body text flow):
  - Paragraphs and runs with bold / italic / underline / strikethrough.
  - Font family and font size.
  - Run text colour.
  - Paragraph alignment (left / center / right / justify).
  - Force-page-break-before-paragraph.
  - Simple single-level bullet and numbered lists (best effort).

What is NOT supported yet (import reports it and recommends against importing
when the unsupported content is significant; export simply cannot emit it):
  - Tables, images / drawings / text boxes, embedded objects.
  - Headers / footers, footnotes / endnotes, comments, equations.
  - Multi-level list numbering definitions, tab stops, styles inheritance,
    sections / columns, fields, tracked changes.

The public API:
  - export_docx(doc, path) -> DocxReport
  - import_docx(path)      -> (Document, DocxReport)

Both return a DocxReport describing what was (not) handled. The editor uses
report.recommend_import / report.recommend_reason to warn the user before an
import, per the product decision: rather than silently dropping unsupported
content, tell the user it is not EDOF-compatible yet and advise against it.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import List, Optional, Tuple, TYPE_CHECKING

if TYPE_CHECKING:
    from edof.format.document import Document

# 1 point = 1/72 inch; 1 inch = 25.4 mm
_MM_PER_PT = 25.4 / 72.0

# Sensible body defaults when a run inherits (EDOF run fields are None=inherit).
_DEFAULT_FONT = "Arial"
_DEFAULT_SIZE_MM = 3.881   # ~= 11 pt, matches the document-mode body style


def _pt_to_mm(pt: float) -> float:
    return float(pt) * _MM_PER_PT


def _mm_to_pt(mm: float) -> float:
    return float(mm) / _MM_PER_PT


# ──────────────────────────────────────────────────────────────────────────
#  Report
# ──────────────────────────────────────────────────────────────────────────
@dataclass
class DocxReport:
    """Outcome of an import or export, surfaced to the user."""
    paragraphs: int = 0
    warnings: List[str] = field(default_factory=list)
    unsupported: List[str] = field(default_factory=list)
    recommend_import: bool = True
    recommend_reason: str = ""

    def summary(self) -> str:
        lines = []
        if self.unsupported:
            lines.append("Not yet EDOF-compatible: " + ", ".join(self.unsupported) + ".")
        if self.warnings:
            lines.extend(self.warnings)
        if not lines:
            lines.append("No compatibility issues found.")
        return "\n".join(lines)


def _require_docx():
    try:
        import docx  # noqa: F401
        return docx
    except Exception as e:  # pragma: no cover - environment dependent
        raise RuntimeError(
            "DOCX support requires the 'python-docx' package "
            "(pip install python-docx)."
        ) from e


# ──────────────────────────────────────────────────────────────────────────
#  Helpers: paragraph source + formatting resolution
# ──────────────────────────────────────────────────────────────────────────
def _source_paragraphs(doc) -> List:
    """Return the EDOF Paragraph list to export.

    Prefers the canonical document-mode body flow. Falls back to gathering
    runs from text boxes across pages so non-document docs still export
    *something* textual (each text box becomes one or more paragraphs)."""
    from edof.format.document_body import Paragraph
    from edof.format.styles import TextRun

    body = getattr(doc, "body", None)
    if body is not None and getattr(body, "paragraphs", None):
        return list(body.paragraphs)

    # Fallback: pull text from any text-bearing objects on the pages.
    paras: List = []
    for page in getattr(doc, "pages", []) or []:
        for obj in getattr(page, "objects", []) or []:
            runs = getattr(obj, "runs", None)
            txt = getattr(obj, "text", None)
            if runs:
                # split the box's runs at newlines into separate paragraphs
                cur: List[TextRun] = []
                for r in runs:
                    segs = (r.text or "").split("\n")
                    for i, seg in enumerate(segs):
                        nr = TextRun(text=seg)
                        for f in ("font_family", "font_size", "bold", "italic",
                                  "underline", "strikethrough", "color",
                                  "alignment"):
                            setattr(nr, f, getattr(r, f, None))
                        cur.append(nr)
                        if i < len(segs) - 1:
                            paras.append(Paragraph(runs=cur))
                            cur = []
                if cur:
                    paras.append(Paragraph(runs=cur))
            elif txt:
                for line in str(txt).split("\n"):
                    paras.append(Paragraph(runs=[TextRun(text=line)]))
    return paras


def _para_alignment(para) -> Optional[str]:
    a = getattr(para, "alignment", None)
    if a:
        return a
    for r in getattr(para, "runs", []) or []:
        if getattr(r, "alignment", None):
            return r.alignment
    return None


def _body_default_line_height(doc) -> float:
    """The document's default line-spacing multiple — taken from the body text
    box style (document mode uses 1.15), so the export matches what the editor
    shows. Falls back to 1.15."""
    for page in getattr(doc, "pages", []) or []:
        for obj in getattr(page, "objects", []) or []:
            nm = (getattr(obj, "name", "") or "").lower()
            st = getattr(obj, "style", None)
            lh = getattr(st, "line_height", None) if st is not None else None
            if lh and ("document_body" in nm or "doc_body" in nm):
                return float(lh)
    # Otherwise take the first text box that carries a line_height.
    for page in getattr(doc, "pages", []) or []:
        for obj in getattr(page, "objects", []) or []:
            st = getattr(obj, "style", None)
            lh = getattr(st, "line_height", None) if st is not None else None
            if lh:
                return float(lh)
    return 1.15


def _para_line_height(para, default_lh: float) -> float:
    """Effective line-spacing multiple for a paragraph: its own override, else
    the max line_height across its runs, else the document default."""
    lh = getattr(para, "line_height", None)
    if lh:
        return float(lh)
    run_lhs = [getattr(r, "line_height", None) for r in (getattr(para, "runs", []) or [])]
    run_lhs = [x for x in run_lhs if x]
    if run_lhs:
        return float(max(run_lhs))
    return float(default_lh)


def _body_default_font_size(doc) -> float:
    """The document's default font size in mm (the body text box style; document
    mode uses 3.881 mm ~= 11 pt). Falls back to the module default."""
    for page in getattr(doc, "pages", []) or []:
        for obj in getattr(page, "objects", []) or []:
            nm = (getattr(obj, "name", "") or "").lower()
            st = getattr(obj, "style", None)
            fs = getattr(st, "font_size", None) if st is not None else None
            if fs and ("document_body" in nm or "doc_body" in nm):
                return float(fs)
    for page in getattr(doc, "pages", []) or []:
        for obj in getattr(page, "objects", []) or []:
            st = getattr(obj, "style", None)
            fs = getattr(st, "font_size", None) if st is not None else None
            if fs:
                return float(fs)
    return _DEFAULT_SIZE_MM


def _para_font_size_mm(para, default_fs: float) -> float:
    """Effective (dominant) font size of a paragraph in mm: the largest run
    size, resolving inherited (None) runs to the document default."""
    sizes = []
    for r in (getattr(para, "runs", []) or []):
        fs = getattr(r, "font_size", None)
        sizes.append(float(fs) if fs else float(default_fs))
    return max(sizes) if sizes else float(default_fs)


# ──────────────────────────────────────────────────────────────────────────
#  EXPORT
# ──────────────────────────────────────────────────────────────────────────
def export_docx(doc, path: str) -> DocxReport:
    """Write an EDOF document to a .docx file. Returns a DocxReport."""
    _require_docx()
    from docx import Document as DocxDocument
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rep = DocxReport()
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    d = DocxDocument()

    # Page geometry: match the EDOF page + margins where we can.
    try:
        sec = d.sections[0]
        pages = getattr(doc, "pages", None) or []
        if pages:
            pw = float(getattr(pages[0], "width", 210.0) or 210.0)
            ph = float(getattr(pages[0], "height", 297.0) or 297.0)
            sec.page_width = Mm(pw)
            sec.page_height = Mm(ph)
        margins = getattr(doc, "margins", None)
        if margins and len(margins) == 4:
            top, right, bottom, left = margins
            sec.top_margin = Mm(float(top))
            sec.right_margin = Mm(float(right))
            sec.bottom_margin = Mm(float(bottom))
            sec.left_margin = Mm(float(left))
    except Exception:
        rep.warnings.append("Page size/margins could not be applied exactly.")

    paras = _source_paragraphs(doc)
    used_list = False
    default_lh = _body_default_line_height(doc)
    default_fs = _body_default_font_size(doc)

    for para in paras:
        runs = getattr(para, "runs", []) or []
        list_kind = getattr(para, "list_kind", "none") or "none"

        p = None
        if list_kind == "bullet":
            try:
                p = d.add_paragraph(style="List Bullet"); used_list = True
            except Exception:
                p = d.add_paragraph()
        elif list_kind == "number":
            try:
                p = d.add_paragraph(style="List Number"); used_list = True
            except Exception:
                p = d.add_paragraph()
        else:
            p = d.add_paragraph()

        al = _para_alignment(para)
        if al in align_map:
            p.alignment = align_map[al]

        # Vertical rhythm: line spacing + space before/after. EDOF's line
        # height is font_size * line_height_multiple (em-based). Word's
        # "multiple" rule instead multiplies the FONT's natural line height
        # (~1.15 em for Arial), which comes out ~15 % taller and made exports
        # spill onto extra pages. So we set an EXACT line height in points equal
        # to EDOF's, which makes Word lay lines out at the same pitch.
        pf = p.paragraph_format
        try:
            lh_mult = _para_line_height(para, default_lh)
            fs_mm = _para_font_size_mm(para, default_fs)
            exact_pt = (fs_mm * lh_mult) / _MM_PER_PT
            pf.line_spacing = Pt(exact_pt)   # Length → EXACTLY rule
        except Exception:
            pass
        try:
            sb = getattr(para, "space_before_mm", None)
            pf.space_before = Mm(float(sb) if sb is not None else 0.0)
        except Exception:
            pass
        try:
            sa = getattr(para, "space_after_mm", None)
            pf.space_after = Mm(float(sa) if sa is not None else 0.0)
        except Exception:
            pass

        try:
            if getattr(para, "page_break_before", False):
                pf.page_break_before = True
        except Exception:
            pass

        if not runs:
            continue

        for r in runs:
            text = r.text or ""
            run = p.add_run(text)
            run.bold = bool(getattr(r, "bold", False))
            run.italic = bool(getattr(r, "italic", False))
            run.underline = bool(getattr(r, "underline", False))
            try:
                if getattr(r, "strikethrough", False):
                    run.font.strike = True
            except Exception:
                pass
            fam = getattr(r, "font_family", None) or _DEFAULT_FONT
            try:
                run.font.name = fam
            except Exception:
                pass
            size_mm = getattr(r, "font_size", None)
            if size_mm is None:
                size_mm = _DEFAULT_SIZE_MM
            try:
                run.font.size = Pt(_mm_to_pt(size_mm))
            except Exception:
                pass
            col = getattr(r, "color", None)
            if col and isinstance(col, (tuple, list)) and len(col) >= 3:
                try:
                    run.font.color.rgb = RGBColor(int(col[0]), int(col[1]), int(col[2]))
                except Exception:
                    pass

    if used_list:
        rep.warnings.append(
            "Lists were exported as single-level Word list styles; "
            "nesting and custom numbering are simplified.")

    rep.paragraphs = len(paras)
    d.save(path)
    return rep


# ──────────────────────────────────────────────────────────────────────────
#  IMPORT — compatibility scan
# ──────────────────────────────────────────────────────────────────────────
def _scan_compat(d, rep: DocxReport) -> None:
    """Populate rep.unsupported / warnings / recommend_* from a docx Document."""
    major: List[str] = []
    minor: List[str] = []

    try:
        if getattr(d, "tables", None):
            major.append("tables")
    except Exception:
        pass
    try:
        if getattr(d, "inline_shapes", None) and len(d.inline_shapes) > 0:
            major.append("images")
    except Exception:
        pass

    # Raw XML probes for things python-docx does not expose directly.
    try:
        body_xml = d.element.body.xml
    except Exception:
        body_xml = ""
    if "w:drawing" in body_xml or "pic:pic" in body_xml or "v:shape" in body_xml:
        if "images" not in major:
            major.append("images / drawings")
    if "wps:txbx" in body_xml or "w:txbxContent" in body_xml:
        major.append("text boxes")
    if "m:oMath" in body_xml:
        major.append("equations")
    if "w:object" in body_xml or "o:OLEObject" in body_xml:
        major.append("embedded objects")

    # Headers / footers with content (minor — body still imports cleanly).
    try:
        for sec in d.sections:
            for hf, label in ((sec.header, "header"), (sec.footer, "footer")):
                if hf is not None and not hf.is_linked_to_previous:
                    if any((p.text or "").strip() for p in hf.paragraphs):
                        minor.append("headers/footers")
                        raise StopIteration
    except StopIteration:
        pass
    except Exception:
        pass

    # Footnotes / comments via package parts.
    try:
        names = [pp.partname for pp in d.part.package.iter_parts()]
        joined = " ".join(str(n) for n in names)
        if "footnotes.xml" in joined or "endnotes.xml" in joined:
            minor.append("footnotes/endnotes")
        if "comments.xml" in joined:
            minor.append("comments")
    except Exception:
        pass

    # De-dup preserving order.
    def _dedup(seq):
        seen = set(); out = []
        for x in seq:
            if x not in seen:
                seen.add(x); out.append(x)
        return out

    major = _dedup(major)
    minor = _dedup(minor)
    rep.unsupported = major + minor

    if major:
        rep.recommend_import = False
        rep.recommend_reason = (
            "This document contains " + ", ".join(major) + ", which EDOF "
            "cannot represent yet. Only the plain text would be imported; "
            "importing is not recommended.")
    elif minor:
        rep.recommend_import = True
        rep.recommend_reason = (
            "This document contains " + ", ".join(minor) + ", which will be "
            "dropped. The text imports fine.")
    else:
        rep.recommend_import = True
        rep.recommend_reason = ""


# ──────────────────────────────────────────────────────────────────────────
#  IMPORT — paragraph reading
# ──────────────────────────────────────────────────────────────────────────
def _docx_align_to_edof(al) -> Optional[str]:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH as W
    except Exception:
        return None
    return {
        W.LEFT: "left", W.CENTER: "center", W.RIGHT: "right",
        W.JUSTIFY: "justify",
    }.get(al, None)


def _detect_list(p) -> Tuple[str, Optional[int]]:
    """Return (list_kind, list_level) for a docx paragraph (best effort)."""
    name = ""
    try:
        name = (p.style.name or "").lower()
    except Exception:
        pass
    kind = "none"
    if "list bullet" in name or "bullet" in name:
        kind = "bullet"
    elif "list number" in name or "number" in name:
        kind = "number"
    else:
        # numbering via pPr/numPr
        try:
            ppr = p._p.pPr
            if ppr is not None and ppr.numPr is not None:
                kind = "number"  # cannot easily tell bullet vs number from numId
        except Exception:
            pass
    if kind == "none":
        return "none", None
    # nesting level
    level = 0
    try:
        ppr = p._p.pPr
        if ppr is not None and ppr.numPr is not None and ppr.numPr.ilvl is not None:
            level = int(ppr.numPr.ilvl.val or 0)
    except Exception:
        level = 0
    return kind, max(0, min(2, level))


def _read_paragraphs(d):
    from edof.format.document_body import Paragraph
    from edof.format.styles import TextRun

    out: List[Paragraph] = []
    for p in d.paragraphs:
        runs_out: List[TextRun] = []
        for r in p.runs:
            txt = r.text or ""
            tr = TextRun(text=txt)
            try: tr.bold = True if r.bold else None
            except Exception: pass
            try: tr.italic = True if r.italic else None
            except Exception: pass
            try: tr.underline = True if r.underline else None
            except Exception: pass
            try:
                if r.font.strike:
                    tr.strikethrough = True
            except Exception:
                pass
            try:
                if r.font.name:
                    tr.font_family = r.font.name
            except Exception:
                pass
            try:
                if r.font.size is not None:
                    tr.font_size = _pt_to_mm(r.font.size.pt)
            except Exception:
                pass
            try:
                rgb = r.font.color.rgb if r.font.color and r.font.color.rgb else None
                if rgb is not None:
                    tr.color = (int(rgb[0]), int(rgb[1]), int(rgb[2]))
            except Exception:
                pass
            runs_out.append(tr)

        if not runs_out:
            runs_out = [TextRun(text=p.text or "")]

        para = Paragraph(runs=runs_out)
        al = None
        try:
            al = _docx_align_to_edof(p.paragraph_format.alignment)
        except Exception:
            al = None
        if al:
            para.alignment = al
        try:
            if p.paragraph_format.page_break_before:
                para.page_break_before = True
        except Exception:
            pass
        # Vertical rhythm back into EDOF.
        try:
            ls = p.paragraph_format.line_spacing
            if isinstance(ls, float):          # MULTIPLE rule → a plain multiplier
                para.line_height = ls
            elif ls is not None:               # EXACTLY/AT_LEAST (a Length in pt)
                # Recover EDOF's em-based multiple: exact_mm / paragraph_font_mm.
                fs_sizes = [tr.font_size for tr in runs_out if getattr(tr, "font_size", None)]
                fs_mm = max(fs_sizes) if fs_sizes else _DEFAULT_SIZE_MM
                if fs_mm > 0:
                    para.line_height = float(ls.mm) / fs_mm
        except Exception:
            pass
        try:
            sb = p.paragraph_format.space_before
            if sb is not None:
                para.space_before_mm = float(sb.mm)
        except Exception:
            pass
        try:
            sa = p.paragraph_format.space_after
            if sa is not None:
                para.space_after_mm = float(sa.mm)
        except Exception:
            pass
        kind, level = _detect_list(p)
        if kind != "none":
            para.list_kind = kind
            para.list_level = level
        out.append(para)

    if not out:
        out = [Paragraph(runs=[TextRun(text="")])]
    return out


def import_docx(path: str) -> Tuple["Document", DocxReport]:
    """Read a .docx into a new EDOF document-mode Document.

    Returns (doc, report). Always returns a usable document containing the
    importable text; the report says what was dropped and whether importing
    is advisable."""
    import edof
    docx = _require_docx()
    from docx import Document as DocxDocument
    from edof.format.document_body import DocumentBody
    from edof.engine.document_paginate import paginate_document

    rep = DocxReport()
    d = DocxDocument(path)

    _scan_compat(d, rep)
    paras = _read_paragraphs(d)
    rep.paragraphs = len(paras)

    # Build a document-mode EDOF doc (A4, 15 mm margins) from the body flow.
    doc = edof.new(width=210.0, height=297.0, title="Imported", dpi=300)
    doc.margins = (15.0, 15.0, 15.0, 15.0)
    doc.mode = "document"
    body = DocumentBody()
    body.page_margins_mm = (15.0, 15.0, 15.0, 15.0)
    body.paragraphs = paras
    doc.body = body

    # Lay the body out across pages. skip_sync keeps body.paragraphs as the
    # source of truth (there are no body text boxes to project back from yet).
    try:
        paginate_document(doc, skip_sync=True)
    except Exception as e:
        rep.warnings.append("Pagination after import failed: %s" % e)

    return doc, rep
